"""
Document Parser Implementation
Handles PDF, DOCX, TXT, and Markdown files with structure preservation
"""

from typing import Dict, Any, Optional, List
import logging
from pathlib import Path
from dataclasses import dataclass

logger = logging.getLogger("rag.processing.document_parser")

# Import parsers with fallback
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class ParsedDocument:
    """Structured representation of parsed document"""
    content: str
    metadata: Dict[str, Any]
    doc_id: str
    filename: str
    file_type: str
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    structure: Optional[Dict[str, Any]] = None  # Headers, sections, etc.


class DocumentParser:
    """
    Multi-format document parser with structure preservation
    Supports: PDF, DOCX, TXT, MD
    """
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md', '.markdown'}
    
    def __init__(self, preserve_structure: bool = True):
        self.preserve_structure = preserve_structure
        
        # Check available parsers
        self.pdf_available = PDF_AVAILABLE
        self.docx_available = DOCX_AVAILABLE
        
        if not self.pdf_available:
            logger.warning("PyPDF2 not available - PDF parsing disabled")
        if not self.docx_available:
            logger.warning("python-docx not available - DOCX parsing disabled")
    
    def parse_file(self, file_path: str, doc_id: Optional[str] = None) -> ParsedDocument:
        """
        Parse document file and return structured content
        
        Args:
            file_path: Path to document file
            doc_id: Optional document ID, defaults to filename
            
        Returns:
            ParsedDocument with content and metadata
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        file_ext = path.suffix.lower()
        if file_ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {file_ext}. Supported: {self.SUPPORTED_EXTENSIONS}")
        
        doc_id = doc_id or path.stem
        filename = path.name
        
        logger.info(f"Parsing {file_ext} document: {filename}")
        
        # Route to appropriate parser
        if file_ext == '.pdf':
            return self._parse_pdf(path, doc_id, filename)
        elif file_ext == '.docx':
            return self._parse_docx(path, doc_id, filename)
        elif file_ext in {'.txt', '.md', '.markdown'}:
            return self._parse_text(path, doc_id, filename, file_ext)
        else:
            raise ValueError(f"No parser available for {file_ext}")
    
    def parse_content(self, 
                     content: str, 
                     filename: str,
                     file_type: str,
                     doc_id: Optional[str] = None) -> ParsedDocument:
        """
        Parse content directly (for uploaded files)
        
        Args:
            content: Raw file content
            filename: Original filename
            file_type: File extension (.pdf, .docx, etc.)
            doc_id: Optional document ID
            
        Returns:
            ParsedDocument with parsed content
        """
        doc_id = doc_id or Path(filename).stem
        
        if file_type == '.txt' or file_type == '.md':
            # Text content can be used directly
            return self._create_parsed_document(
                content=content,
                doc_id=doc_id,
                filename=filename,
                file_type=file_type,
                metadata={"parsing_method": "direct_text"}
            )
        else:
            # Binary formats need file-based parsing
            raise ValueError(f"Content parsing not supported for {file_type}, use file path instead")
    
    def _parse_pdf(self, path: Path, doc_id: str, filename: str) -> ParsedDocument:
        """Parse PDF file using PyPDF2"""
        if not self.pdf_available:
            raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
        
        try:
            content_parts = []
            metadata = {"parsing_method": "PyPDF2"}
            structure = {"pages": []} if self.preserve_structure else None
            
            with open(path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                # Extract metadata
                if pdf_reader.metadata:
                    pdf_meta = pdf_reader.metadata
                    metadata.update({
                        "title": getattr(pdf_meta, 'title', None),
                        "author": getattr(pdf_meta, 'author', None),
                        "subject": getattr(pdf_meta, 'subject', None),
                        "creator": getattr(pdf_meta, 'creator', None),
                        "producer": getattr(pdf_meta, 'producer', None),
                        "creation_date": str(getattr(pdf_meta, 'creation_date', None)),
                        "modification_date": str(getattr(pdf_meta, 'modification_date', None))
                    })
                
                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            content_parts.append(f"\n--- Page {page_num + 1} ---\n")
                            content_parts.append(page_text)
                            
                            if structure:
                                structure["pages"].append({
                                    "page_number": page_num + 1,
                                    "char_start": len("".join(content_parts[:-2])),
                                    "char_end": len("".join(content_parts)),
                                    "text_length": len(page_text)
                                })
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
                
                content = "".join(content_parts)
                
                return self._create_parsed_document(
                    content=content,
                    doc_id=doc_id,
                    filename=filename,
                    file_type='.pdf',
                    metadata=metadata,
                    page_count=page_count,
                    structure=structure
                )
                
        except Exception as e:
            logger.error(f"PDF parsing failed for {filename}: {e}")
            raise
    
    def _parse_docx(self, path: Path, doc_id: str, filename: str) -> ParsedDocument:
        """Parse DOCX file using python-docx"""
        if not self.docx_available:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        try:
            doc = DocxDocument(str(path))
            content_parts = []
            metadata = {"parsing_method": "python-docx"}
            structure = {"paragraphs": [], "headers": []} if self.preserve_structure else None
            
            # Extract core properties
            core_props = doc.core_properties
            if core_props:
                metadata.update({
                    "title": core_props.title,
                    "author": core_props.author,
                    "subject": core_props.subject,
                    "keywords": core_props.keywords,
                    "comments": core_props.comments,
                    "created": str(core_props.created),
                    "modified": str(core_props.modified),
                    "revision": core_props.revision
                })
            
            # Extract text from paragraphs
            for para_num, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                if para_text:
                    # Check if it's a heading
                    is_header = paragraph.style.name.startswith('Heading')
                    
                    if is_header:
                        content_parts.append(f"\n\n## {para_text}\n")
                        if structure:
                            structure["headers"].append({
                                "text": para_text,
                                "level": paragraph.style.name,
                                "paragraph_index": para_num
                            })
                    else:
                        content_parts.append(para_text + "\n")
                    
                    if structure:
                        structure["paragraphs"].append({
                            "index": para_num,
                            "text": para_text,
                            "is_header": is_header,
                            "style": paragraph.style.name
                        })
            
            content = "".join(content_parts)
            
            return self._create_parsed_document(
                content=content,
                doc_id=doc_id,
                filename=filename,
                file_type='.docx',
                metadata=metadata,
                structure=structure
            )
            
        except Exception as e:
            logger.error(f"DOCX parsing failed for {filename}: {e}")
            raise
    
    def _parse_text(self, path: Path, doc_id: str, filename: str, file_ext: str) -> ParsedDocument:
        """Parse text/markdown files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(path, 'r', encoding=encoding) as file:
                        content = file.read()
                        used_encoding = encoding
                        break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError(f"Could not decode file {filename} with any supported encoding")
            
            metadata = {
                "parsing_method": "text",
                "encoding": used_encoding,
                "file_size": path.stat().st_size
            }
            
            # Basic structure detection for markdown
            structure = None
            if file_ext in {'.md', '.markdown'} and self.preserve_structure:
                structure = self._detect_markdown_structure(content)
            
            return self._create_parsed_document(
                content=content,
                doc_id=doc_id,
                filename=filename,
                file_type=file_ext,
                metadata=metadata,
                structure=structure
            )
            
        except Exception as e:
            logger.error(f"Text parsing failed for {filename}: {e}")
            raise
    
    def _detect_markdown_structure(self, content: str) -> Dict[str, Any]:
        """Detect markdown structure (headers, lists, code blocks)"""
        import re
        
        structure = {
            "headers": [],
            "code_blocks": [],
            "lists": []
        }
        
        lines = content.split('\n')
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            # Headers
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line_stripped)
            if header_match:
                level = len(header_match.group(1))
                text = header_match.group(2)
                structure["headers"].append({
                    "level": level,
                    "text": text,
                    "line_number": i + 1
                })
            
            # Code blocks
            elif line_stripped.startswith('```'):
                structure["code_blocks"].append({
                    "line_number": i + 1,
                    "language": line_stripped[3:].strip()
                })
            
            # Lists
            elif re.match(r'^[\s]*[-*+]\s+|^\s*\d+\.\s+', line_stripped):
                structure["lists"].append({
                    "line_number": i + 1,
                    "type": "ordered" if re.match(r'^\s*\d+\.', line_stripped) else "unordered",
                    "text": line_stripped
                })
        
        return structure
    
    def _create_parsed_document(self,
                               content: str,
                               doc_id: str,
                               filename: str,
                               file_type: str,
                               metadata: Dict[str, Any],
                               page_count: Optional[int] = None,
                               structure: Optional[Dict[str, Any]] = None) -> ParsedDocument:
        """Create ParsedDocument with calculated metrics"""
        
        # Calculate word count
        word_count = len(content.split()) if content else 0
        
        # Add parsing metadata
        metadata.update({
            "parsed_at": str(__import__('datetime').datetime.now()),
            "word_count": word_count,
            "char_count": len(content),
            "preserve_structure": self.preserve_structure
        })
        
        return ParsedDocument(
            content=content,
            metadata=metadata,
            doc_id=doc_id,
            filename=filename,
            file_type=file_type,
            page_count=page_count,
            word_count=word_count,
            structure=structure
        )
    
    def get_supported_types(self) -> Dict[str, bool]:
        """Get supported file types and their availability"""
        return {
            '.txt': True,
            '.md': True,
            '.markdown': True,
            '.pdf': self.pdf_available,
            '.docx': self.docx_available
        }